import os
import sqlite3
import json
import time
import asyncio
import traceback
import logging
import csv
from functools import lru_cache
from pathlib import Path
from typing import List, Dict, Any
from fastapi import FastAPI, UploadFile, File, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pdfplumber
import google.generativeai as genai
from datetime import datetime
import io

# ================= 日志配置 =================
logging.basicConfig(
    level=logging.INFO,  # 生产环境使用 INFO，调试时改为 DEBUG
    format='%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)
# =========================================

# ================= 配置 =================
API_KEY = os.environ.get("API_KEY") 
if not API_KEY:
    logger.warning("API_KEY not found in environment variables! Please set it.")
MODEL_NAME = 'gemini-2.5-pro'

# Adapt DB path for Render persistent storage
if os.environ.get('RENDER'):
    # Define paths for both V1 and V2 databases
    DB_PATH = "/var/lib/data/knowledge_base.db"
    repo_db_path = "knowledge_base.db"
    
    DB_V2_PATH = "/var/lib/data/knowledge_base_v2.db"
    repo_db_v2_path = "knowledge_base_v2.db"
    
    # --- Seeding Logic for V1 DB ---
    should_seed = False
    if not os.path.exists(DB_PATH):
        should_seed = True
        logger.info("V1 DB not found on disk. Marking for seed.")
    else:
        # Check if existing DB is valid (has tables)
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            cursor.execute("SELECT count(*) FROM sqlite_master WHERE type='table' AND name='knowledge_points'")
            if cursor.fetchone()[0] == 0:
                logger.warning("Existing V1 DB on disk is missing 'knowledge_points' table. Marking for re-seed...")
                should_seed = True
            conn.close()
        except Exception as e:
             logger.warning(f"Existing V1 DB check failed: {e}. Marking for re-seed...")
             should_seed = True

    if should_seed:
        if os.path.exists(repo_db_path):
            try:
                import shutil
                logger.info(f"Seeding V1 database from {repo_db_path} to {DB_PATH}...")
                os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
                shutil.copy(repo_db_path, DB_PATH)
                logger.info("V1 Database seeded successfully.")
            except Exception as e:
                logger.error(f"Failed to seed V1 database: {e}")
        else:
            logger.warning(f"No seed V1 database found at {repo_db_path}. Starting with empty DB.")
            
    # --- Seeding Logic for V2 DB (Precision Search) ---
    should_seed_v2 = False
    if not os.path.exists(DB_V2_PATH):
        should_seed_v2 = True
        logger.info("V2 DB not found on disk. Marking for seed.")
    else:
        try:
            conn = sqlite3.connect(DB_V2_PATH)
            cursor = conn.cursor()
            cursor.execute("SELECT count(*) FROM sqlite_master WHERE type='table' AND name='theories'")
            if cursor.fetchone()[0] == 0:
                logger.warning("Existing V2 DB on disk is missing 'theories' table. Marking for re-seed...")
                should_seed_v2 = True
            conn.close()
        except Exception as e:
            logger.warning(f"Existing V2 DB check failed: {e}. Marking for re-seed...")
            should_seed_v2 = True
            
    if should_seed_v2:
        if os.path.exists(repo_db_v2_path):
            try:
                import shutil
                logger.info(f"Seeding V2 database from {repo_db_v2_path} to {DB_V2_PATH}...")
                os.makedirs(os.path.dirname(DB_V2_PATH), exist_ok=True)
                shutil.copy(repo_db_v2_path, DB_V2_PATH)
                logger.info("V2 Database seeded successfully.")
            except Exception as e:
                logger.error(f"Failed to seed V2 database: {e}")
        else:
            logger.warning(f"No seed V2 database found at {repo_db_v2_path}. Starting with empty DB.")

else:
    DB_PATH = "knowledge_base.db"
    # Local path for V2
    DB_V2_PATH = "knowledge_base_v2.db"

# Proxy settings: Only use proxy if NOT on Render (Render is in US, no proxy needed)
if os.environ.get('RENDER'):
    PROXY_URL = None
    # Clear proxy env vars if they exist
    os.environ.pop('HTTP_PROXY', None)
    os.environ.pop('HTTPS_PROXY', None)
    os.environ.pop('NO_PROXY', None)
else:
    PROXY_URL = "http://127.0.0.1:7897" 
    # 环境变量配置
    os.environ['HTTP_PROXY'] = PROXY_URL
    os.environ['HTTPS_PROXY'] = PROXY_URL
    os.environ['NO_PROXY'] = "localhost,127.0.0.1"

# =======================================

# --- 启动自检 ---
logger.info("="*60)
logger.info(">>> FDC Case Check System Starting Up...")
logger.info(f">>> Model: {MODEL_NAME}")
logger.info(f">>> Database: {DB_PATH}")
logger.info(f">>> Proxy: {PROXY_URL}")

try:
    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel(MODEL_NAME)
    logger.info(">>> Gemini API configured successfully")
except Exception as e:
    logger.error(f">>> Gemini configuration failed: {e}", exc_info=True)

# 检查数据库
try:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM knowledge_points")
    count = c.fetchone()[0]
    conn.close()
    logger.info(f">>> Database loaded: {count} knowledge points")
except Exception as e:
    logger.error(f">>> Database check failed: {e}", exc_info=True)

logger.info("="*60)

app = FastAPI(title="FDC Case Check API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- 数据模型 ---
class ReportData(BaseModel):
    meta: Dict[str, Any]
    summary: Dict[str, Any]
    results: Dict[str, Any]
    report_type: str = "dashboard" # New field for report type

# --- 辅助函数 ---

async def send_heartbeat(websocket: WebSocket, interval: int = 10):
    """WebSocket 心跳任务，定期发送ping保持连接活跃"""
    try:
        while True:
            await asyncio.sleep(interval)
            try:
                await websocket.send_json({"type": "heartbeat", "timestamp": datetime.now().isoformat()})
            except Exception as e:
                logger.warning(f"Heartbeat failed: {e}")
                break
    except asyncio.CancelledError:
        logger.info("Heartbeat task cancelled")

async def retry_async_operation(func, max_retries=3, delay=2, *args, **kwargs):
    """带重试机制的异步操作"""
    last_error = None
    for attempt in range(max_retries):
        try:
            logger.info(f"Attempt {attempt + 1}/{max_retries} for {func.__name__}")
            return await func(*args, **kwargs)
        except Exception as e:
            last_error = e
            logger.warning(f"Attempt {attempt + 1} failed: {e}")
            if attempt < max_retries - 1:
                await asyncio.sleep(delay * (attempt + 1))  # 指数退避
            else:
                logger.error(f"All {max_retries} attempts failed for {func.__name__}: {e}")
    raise last_error

# --- 核心工具函数 ---

def extract_text_from_pdf_bytes(file_bytes):
    """从PDF字节流提取文本，带文件大小限制"""
    import io

    # 限制文件大小为50MB
    MAX_FILE_SIZE = 50 * 1024 * 1024
    if len(file_bytes) > MAX_FILE_SIZE:
        logger.error(f"PDF file too large: {len(file_bytes) / 1024 / 1024:.2f} MB (max: 50 MB)")
        return None

    text = ""
    try:
        logger.info(f"Extracting text from PDF ({len(file_bytes) / 1024:.2f} KB)")
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t:
                    text += t + "\n"
                # Log progress every 10 pages OR on the very last page
                if (i + 1) % 10 == 0 or (i + 1) == len(pdf.pages):
                    logger.info(f"Processed {i + 1}/{len(pdf.pages)} pages")
        logger.info(f"Successfully extracted {len(text)} characters")
    except Exception as e:
        logger.error(f"Error reading PDF: {e}", exc_info=True)
        return None
    return text

async def extract_new_concepts(content, websocket=None):
    """使用AI提取理论框架（不再提取知识点）"""
    truncated_content = content[:50000]
    logger.info(f"Starting AI extraction (content length: {len(truncated_content)} chars)")
    logger.debug(f"Text preview: {truncated_content[:200]}")

    prompt = f"""
你是一个商业案例分析专家。请仔细分析以下商业案例文本，**全面提取**案例中涉及的理论框架和管理模型。

核心原则：**LOGIC & EXPLICIT（逻辑与显性结合）** - 既要提取明确提及的理论，也要提取**逻辑显包含**的理论。

提取标准：
1. **显性提及 (Explicit)**：文中直接出现了理论名称（如"SWOT"、"波特五力"）。
2. **隐性逻辑 (Implicit)**：文中虽然没有直接说出理论名称，但其论述逻辑、分析维度完全符合某经典理论的定义。
   - *例如*：文中分析了"供应商议价能力"和"潜在进入者"，即使没说"波特五力"，也要提取"波特五力模型"。

提取范围：
- 经典商业理论（波特系列、营销4P/4C、STP等）
- 战略管理工具（BCG矩阵、价值链、商业模式画布等）
- 创新与创业理论（精益创业、破坏性创新、跨越鸿沟等）
- 组织与管理模型（OKR、KPI、阿米巴经营等）

**注意**：
- 只要文中涉及了该理论的核心要素或分析逻辑，**就算**。
- context 字段必须直接引用原文，不能是 AI 生成的摘要。

输出格式（严格JSON）：
{{
  "meta": {{
    "title": "案例标题（从文中提取）"
  }},
  "items": [
    {{
      "category": "理论",
      "standard_name": "波特五力模型",
      "mention_type": "explicit",  // 必须是 "explicit" 或 "implicit"
      "context": "文中出现的上下文（必须是原文引用，禁止概括总结，最多150字）...",
      "rationale": "AI标注：为什么这段语境符合该理论（简练解释，不超过50字）"
    }}
  ]
}}

案例文本：
{truncated_content}

请输出符合上述格式的JSON（不要包含markdown代码块标记）。
    """

    async def _extract():
        try:
            model = genai.GenerativeModel(MODEL_NAME)
            logger.info("Sending request to Gemini API...")
            start_time = time.time()

            res = await asyncio.to_thread(
                model.generate_content,
                prompt,
                generation_config={"temperature": 0.4},
                request_options={'timeout': 300}  # 缩短到5分钟
            )

            elapsed = time.time() - start_time
            logger.info(f"Gemini API responded in {elapsed:.2f}s")

            raw = res.text.strip()
            logger.debug(f"Raw AI response (first 500 chars): {raw[:500]}")

            # 清理 markdown 代码块
            if raw.startswith("```"):
                # 移除开头的 ```json 或 ```
                lines = raw.split('\n')
                raw = '\n'.join(lines[1:])  # 跳过第一行
                if raw.endswith("```"):
                    raw = raw[:-3]  # 移除结尾的 ```

            # 提取 JSON 对象
            start = raw.find('{')
            end = raw.rfind('}') + 1
            if start != -1 and end > start:
                raw = raw[start:end]

            logger.debug(f"Cleaned JSON (first 500 chars): {raw[:500]}")

            try:
                result = json.loads(raw)
            except json.JSONDecodeError as e:
                logger.error(f"JSON parse error: {e}")
                logger.error(f"Problematic JSON: {raw}")
                raise

            logger.info(f"Successfully extracted {len(result.get('items', []))} items")

            # 如果提取为空，记录完整响应
            if not result.get('items'):
                logger.warning(f"AI returned empty items! Full response: {raw}")

            return result
        except Exception as e:
            logger.error(f"Extraction attempt failed: {e}", exc_info=True)
            raise

    # 使用重试机制
    return await retry_async_operation(_extract, max_retries=3, delay=3)

def get_db_tags(category, db_path=DB_PATH, table_name="knowledge_points"):
    """从数据库获取标签，带错误处理"""
    conn = None
    try:
        logger.info(f"Fetching DB tags for category: {category} from {db_path} ({table_name})")
        conn = sqlite3.connect(db_path, timeout=30)
        c = conn.cursor()
        
        if category:
            c.execute(f"SELECT DISTINCT standard_name FROM {table_name} WHERE category=?", (category,))
        else:
            c.execute(f"SELECT DISTINCT standard_name FROM {table_name}")

        tags = [row[0] for row in c.fetchall()]
        logger.info(f"Found {len(tags)} tags")
        return tags
    except Exception as e:
        logger.error(f"Error fetching DB tags for {category}: {e}", exc_info=True)
        return []
    finally:
        if conn:
            conn.close()

async def ai_strict_mapping(new_items, db_tags, category_name):
    """AI三级匹配，支持重试，分批处理大数据集"""
    if not new_items or not db_tags:
        logger.warning(f"Skipping mapping for {category_name}: empty input")
        return {}

    new_terms_list = list(set([item['standard_name'] for item in new_items]))
    total_new_terms = len(new_terms_list)
    total_db_tags = len(db_tags)

    logger.info(f"Mapping {total_new_terms} terms for {category_name} against {total_db_tags} DB tags")

    # 策略：如果数据量太大，分批处理
    # 理论部分：少量术语，不分批
    # 知识点部分：大量术语，必须分批
    if category_name in ['理论', 'Direct Theory', 'All Theories']:
        MAX_NEW_TERMS_PER_BATCH = 100  # 理论允许更多
        MAX_DB_TAGS = 5000
    else:  # 知识点
        MAX_NEW_TERMS_PER_BATCH = 30   # 知识点严格限制，避免 JSON 过大
        MAX_DB_TAGS = 2000              # 限制数据库标签数量

    # 如果数据库标签太多，只使用前N个（简化处理）
    if total_db_tags > MAX_DB_TAGS:
        logger.warning(f"DB tags too large ({total_db_tags}), using first {MAX_DB_TAGS} tags")
        db_tags = db_tags[:MAX_DB_TAGS]

    # 如果新术语太多，分批处理
    if total_new_terms > MAX_NEW_TERMS_PER_BATCH:
        logger.info(f"Splitting {total_new_terms} new terms into batches of {MAX_NEW_TERMS_PER_BATCH}")

        all_mappings = {}
        batch_count = (total_new_terms + MAX_NEW_TERMS_PER_BATCH - 1) // MAX_NEW_TERMS_PER_BATCH

        for i in range(batch_count):
            start_idx = i * MAX_NEW_TERMS_PER_BATCH
            end_idx = min((i + 1) * MAX_NEW_TERMS_PER_BATCH, total_new_terms)
            batch_terms = new_terms_list[start_idx:end_idx]

            logger.info(f"Processing batch {i+1}/{batch_count} ({len(batch_terms)} terms)")

            # 处理这批术语
            batch_mapping = await _map_batch(batch_terms, db_tags, category_name)

            # 合并结果
            all_mappings.update(batch_mapping)

        logger.info(f"Completed all {batch_count} batches, total {len(all_mappings)} matches")
        return all_mappings
    else:
        # 数据量不大，直接处理
        return await _map_batch(new_terms_list, db_tags, category_name)

async def _map_batch(new_terms_list, db_tags, category_name):
    """处理单批术语的匹配"""

    prompt = f"""
你是一个精准语义匹配专家。请将新文档中的术语与历史数据库中的标准术语进行**精准三级匹配**。

核心原则：**只匹配相同的概念/理论，根据表述精确度分级**

输入：
1. 新文档术语列表: {json.dumps(new_terms_list, ensure_ascii=False)}
2. 历史数据库标准术语列表: {json.dumps(db_tags, ensure_ascii=False)}

匹配规则（必须是同一概念）：

- **高度重合**（match_level: "high"）：完全相同的概念，不同表述方式
  ✅ 正确示例：
    "SWOT分析" ≈ "SWOT" ≈ "SWOT分析法" ≈ "SWOT Analysis"
    "波特五力模型" ≈ "波特五力" ≈ "五力模型" ≈ "Porter's Five Forces"
    "市场细分" ≈ "目标市场细分" ≈ "市场分割" ≈ "Market Segmentation"
  ❌ 错误示例：
    "SWOT分析" ≠ "PEST分析"（不同理论，不能匹配）
    "市场细分" ≠ "市场调研"（不同概念，不能匹配）

- **次重合**（match_level: "medium"）：同一概念的细分或特定应用
  ✅ 正确示例：
    "SWOT分析" ≈ "内部优势分析（SWOT的一部分）"
    "价值链分析" ≈ "价值链优化（价值链的应用）"
    "品牌定位" ≈ "品牌市场定位（品牌定位的具体化）"

- **重合**（match_level: "low"）：同一概念的变体、扩展或相关衍生
  ✅ 正确示例：
    "SWOT分析" ≈ "TOWS矩阵（SWOT的逆向应用）"
    "波特五力" ≈ "竞争五力扩展模型（波特五力的扩展）"
    "传统4P" ≈ "7P营销组合（4P的扩展）"

输出格式（严格JSON，**不要超过1000行**）：
{{
  "DB标准术语1": {{
    "matched_terms": ["新术语1", "新术语2"],
    "match_level": "high"
  }},
  "DB标准术语2": {{
    "matched_terms": ["新术语3"],
    "match_level": "medium"
  }}
}}

**严格规则**：
1. **必须是同一概念/理论才能匹配**（不同概念绝对不匹配）
2. 同一概念根据表述差异分为三级：
   - high：完全相同的概念，不同说法（同义词、翻译、简写）
   - medium：同一概念的子集或特定应用
   - low：同一概念的变体或扩展版本
3. 如果新术语与数据库中任何术语都不是同一概念，则不匹配
4. **输出纯JSON，不要包含markdown代码块**
5. **优先匹配高度重合，确保精准度**

请输出匹配结果：
    """

    async def _do_map():
        try:
            model = genai.GenerativeModel(MODEL_NAME)
            logger.info(f"Starting AI mapping for {category_name} (batch: {len(new_terms_list)} terms)...")
            start_time = time.time()

            res = await asyncio.to_thread(
                model.generate_content,
                prompt,
                generation_config={"temperature": 0.1},
                request_options={'timeout': 300}  # 5分钟超时
            )

            elapsed = time.time() - start_time
            logger.info(f"Mapping completed in {elapsed:.2f}s for {category_name}")

            raw = res.text.strip()

            # 清理 markdown 代码块
            if raw.startswith("```"):
                lines = raw.split('\n')
                raw = '\n'.join(lines[1:])
                if raw.endswith("```"):
                    raw = raw[:-3]

            # 提取 JSON
            start = raw.find('{')
            end = raw.rfind('}') + 1
            if start != -1 and end > start:
                raw = raw[start:end]

            # 解析 JSON
            try:
                result = json.loads(raw)
            except json.JSONDecodeError as e:
                logger.error(f"JSON parse error: {e}")
                logger.error(f"Problematic JSON (first 1000 chars): {raw[:1000]}")
                # 如果 JSON 太大或格式错误，返回空结果而不是崩溃
                logger.warning("Returning empty mapping due to JSON parse error")
                return {}

            logger.info(f"Mapped {len(result)} matches for {category_name}")
            return result
        except Exception as e:
            logger.error(f"Mapping attempt failed for {category_name}: {e}", exc_info=True)
            raise

    return await retry_async_operation(_do_map, max_retries=2, delay=5)  # 减少重试次数

async def consolidate_results(results, category_name):
    """
    使用 Gemini 2.5 Pro 归一化合并相同概念，并进行严格去重
    """
    if not results or len(results) <= 1:
        logger.info(f"No consolidation needed for {category_name} ({len(results)} items)")
        return results

    logger.info(f"Consolidating {len(results)} {category_name} results...")

    # 准备数据给 AI 做归一化映射
    terms_data = []
    for r in results:
        terms_data.append({
            "db_term": r['db_term'],
            "matched_terms": r.get('matched_terms', []),
            "match_level": r.get('match_level', 'high')
        })

    prompt = f"""
    你是一个概念归一化专家。请将以下{category_name}列表中**表示同一概念的不同表述**合并成统一条目。

    输入数据：
    {json.dumps(terms_data, ensure_ascii=False, indent=2)}

    任务：
    1. 识别哪些术语表示同一个概念。
    2. 为每组同一概念选择一个**标准中文名称**。
    3. **去除冗余英文**：不要在中文后加括号附带英文原名。
    4. **保留通用缩写**：如果中文惯用语包含缩写（如SWOT分析），请保留，不要强行翻译（不要写成态势分析法）。

    输出格式（JSON数组）：
    [
      {{
        "standard_term": "资源基础理论",
        "variants": ["资源基础观", "资源基础理论", "RBV", "Resource-Based View"]
      }},
      {{
        "standard_term": "SWOT分析",
        "variants": ["SWOT", "SWOT Analysis", "态势分析法"]
      }}
    ]

    **规则**：
    1. 只合并**完全相同的概念**。
    2. 优先选择最通用的中文学术称谓。
    3. 如果不确定是否同一概念，保持独立，不要输出。

    请输出JSON数组（不要markdown代码块）：
    """

    consolidation_map = []
    try:
        # 使用 Gemini 2.5 Pro
        model = genai.GenerativeModel(MODEL_NAME)
        logger.info(f"Calling AI for consolidation map...")
        
        res = await asyncio.to_thread(
            model.generate_content,
            prompt,
            generation_config={"temperature": 0.1},
            request_options={'timeout': 120}
        )

        raw = res.text.strip()
        if raw.startswith("```"):
            lines = raw.split('\n')
            raw = '\n'.join(lines[1:])
            if raw.endswith("```"):
                raw = raw[:-3]
        
        start = raw.find('[')
        end = raw.rfind(']') + 1
        if start != -1 and end > start:
            raw = raw[start:end]
            consolidation_map = json.loads(raw)
            
    except Exception as e:
        logger.error(f"AI Consolidation mapping failed: {e}", exc_info=True)
        # 继续执行，依靠代码去重逻辑

    # --- 严格去重与合并逻辑 ---
    
    # 1. 建立 变体 -> 标准词 的映射字典
    variant_to_standard = {}
    for group in consolidation_map:
        std = group['standard_term']
        for v in group['variants']:
            variant_to_standard[v] = std

    # 2. 聚合结果
    final_results_map = {}

    for r in results:
        # 确定标准名称：AI映射 > 原名
        original_term = r['db_term']
        standard_term = variant_to_standard.get(original_term, original_term)
        
        if standard_term not in final_results_map:
            # 初始化新条目
            final_results_map[standard_term] = {
                "db_term": standard_term,
                "matched_terms": set(r.get('matched_terms', [])),
                "new_contexts": [r.get('new_context', '')],
                "rationales": [r.get('rationale', '')],  # Collect rationales
                "evidence": r.get('evidence', [])
            }
        else:
            # 合并到现有条目
            entry = final_results_map[standard_term]
            entry['matched_terms'].update(r.get('matched_terms', []))
            if r.get('new_context'):
                entry['new_contexts'].append(r['new_context'])
            if r.get('rationale'):
                entry['rationales'].append(r['rationale'])
            
            # 合并证据（去重）
            existing_keys = set((e['filename'], e['year']) for e in entry['evidence'])
            for ev in r.get('evidence', []):
                key = (ev['filename'], ev['year'])
                if key not in existing_keys:
                    entry['evidence'].append(ev)
                    existing_keys.add(key)

    # 3. 转换为列表并计算最终属性
    consolidated_results = []
    for item in final_results_map.values():
        # 选择最长的上下文
        best_context = ""
        if item['new_contexts']:
            best_context = max(item['new_contexts'], key=len)
        
        # 选择最长的 rationale
        best_rationale = ""
        if item['rationales']:
            best_rationale = max(item['rationales'], key=len)
            
        # 计算风险等级
        count = len(item['evidence'])
        if count >= 5:
            match_level = 'high'
            risk_level = '高度重合'
        elif count >= 3:
            match_level = 'medium'
            risk_level = '次重合'
        else:
            match_level = 'low'
            risk_level = '重合'

        consolidated_results.append({
            "risk_level": risk_level,
            "db_term": item['db_term'],
            "matched_terms": list(item['matched_terms']),
            "new_context": best_context,
            "rationale": best_rationale, # Add Rationale
            "evidence": item['evidence'],
            "match_level": match_level
        })

    # 4. 排序：高度重合 > 次重合 > 重合
    risk_priority = {"高度重合": 0, "次重合": 1, "重合": 2}
    consolidated_results.sort(key=lambda x: risk_priority.get(x.get('risk_level', '高度重合'), 99))

    logger.info(f"Consolidated {len(results)} → {len(consolidated_results)} {category_name}")
    return consolidated_results

async def translate_to_chinese_academic(results_list):
    """
    将结果列表中的 db_term 翻译为标准的中文学术名称。
    使用 Gemini 2.5 Pro 进行翻译。
    """
    if not results_list:
        return results_list

    # 1. 提取唯一的 db_term
    unique_terms = list(set(r['db_term'] for r in results_list))
    if not unique_terms:
        return results_list
        
    logger.info(f"Translating {len(unique_terms)} terms to Chinese...")

    prompt = f"""
    你是一个商业管理理论翻译专家。请将以下术语列表翻译为**标准的简体中文学术名称**。
    
    规则：
    1. **去除英文翻译/备注**：不要包含括号内的英文全称或解释。
    2. **保留通用缩写**：如果该理论在中文学术界通用惯例中包含英文缩写（如 "SWOT分析"、"PEST分析"、"4P营销理论"），请**保留该缩写**，不要强行翻译成生僻中文（如不要把SWOT强行翻成态势分析法）。
    3. 如果是纯英文术语，请翻译成最通用的中文标准称谓。
    4. 输出严格的 JSON 字典。
    
    示例：
    - "SWOT Analysis" -> "SWOT分析" (保留SWOT)
    - "Porter's Five Forces (P5F)" -> "波特五力模型" (去除P5F)
    - "Resource-Based View" -> "资源基础理论"
    - "PESTEL" -> "PESTEL模型"
    
    输入列表：
    {json.dumps(unique_terms, ensure_ascii=False, indent=2)}
    
    输出 JSON 格式：
    {{
      "Original Term 1": "标准中文名称1",
      "Original Term 2": "标准中文名称2"
    }}
    """

    try:
        # 强制使用 gemini-2.5-pro
        model = genai.GenerativeModel('gemini-2.5-pro')
        
        # 运行生成
        res = await asyncio.to_thread(
            model.generate_content,
            prompt,
            generation_config={"temperature": 0.1},
            request_options={'timeout': 60}
        )
        
        raw = res.text.strip()
        # 清理 markdown
        if raw.startswith("```"):
            lines = raw.split('\n')
            raw = '\n'.join(lines[1:])
            if raw.endswith("```"):
                raw = raw[:-3]
        
        # 提取 JSON
        start = raw.find('{')
        end = raw.rfind('}') + 1
        if start != -1 and end > start:
            raw = raw[start:end]
            translation_map = json.loads(raw)
            
            # 更新结果
            for item in results_list:
                original = item['db_term']
                if original in translation_map:
                    item['db_term'] = translation_map[original]
            
            logger.info("Translation completed successfully.")
            
    except Exception as e:
        logger.error(f"Translation failed: {e}")
        # 失败时保持原样
        
    return results_list

def generate_analysis_summary(explicit_results, implicit_results, all_new_items):
    """
    生成智能分析总结 - 区分显性与隐性
    """
    try:
        # 1. 基础统计
        total_items = len(all_new_items)
        
        # 显性统计
        exp_high = sum(1 for r in explicit_results if r['match_level'] == 'high')
        exp_med = sum(1 for r in explicit_results if r['match_level'] == 'medium')
        exp_low = sum(1 for r in explicit_results if r['match_level'] == 'low')
        total_explicit_matches = exp_high + exp_med + exp_low

        # 隐性统计
        imp_high = sum(1 for r in implicit_results if r['match_level'] == 'high')
        imp_med = sum(1 for r in implicit_results if r['match_level'] == 'medium')
        imp_low = sum(1 for r in implicit_results if r['match_level'] == 'low')
        total_implicit_matches = imp_high + imp_med + imp_low

        total_matches = total_explicit_matches + total_implicit_matches

        # 计算重复率 - 已移除显示，设为0
        overlap_rate = 0

        # 2. 风险等级综合评估 - 只看显性理论匹配
        # 0-2为低风险 3-4为中等风险 5及5以上为高风险
        if total_explicit_matches >= 5:
            risk_level = "高风险"
            risk_color = "red"
            risk_desc = f"检测到 {total_explicit_matches} 处显性理论匹配，风险较高，建议大幅修改。"
        elif total_explicit_matches >= 3:
            risk_level = "中等风险"
            risk_color = "orange"
            risk_desc = f"检测到 {total_explicit_matches} 处显性理论匹配，存在一定风险，建议优化。"
        else:
            risk_level = "低风险"
            risk_color = "green"
            risk_desc = f"仅检测到 {total_explicit_matches} 处显性理论匹配，原创度良好。"

        # 提取高频理论名称
        high_overlap_theories = [r['db_term'] for r in explicit_results if r['match_level'] == 'high']

        # 3. 历史案例统计
        all_evidence = []
        for r in explicit_results + implicit_results:
            all_evidence.extend(r.get('evidence', []))

        year_stats = {}
        case_stats = {}

        for ev in all_evidence:
            year = ev.get('year', 'Unknown')
            fname = ev.get('filename', 'Unknown')
            year_stats[year] = year_stats.get(year, 0) + 1
            
            if fname not in case_stats:
                case_stats[fname] = {"filename": fname, "year": year, "count": 0}
            case_stats[fname]["count"] += 1
        
        sorted_cases = sorted(case_stats.values(), key=lambda x: x['count'], reverse=True)
        top_cases = sorted_cases[:10]
        top_years = sorted(year_stats.items(), key=lambda x: x[1], reverse=True)[:5]

        # 4. 生成建议
        recommendations = []
        if total_explicit_matches >= 5:
            recommendations.append({
                "level": "critical",
                "title": "显性框架重写",
                "description": f"文中直接使用了 {total_explicit_matches} 个理论框架。",
                "action": "建议完全替换为更具针对性的行业模型，或结合企业特性进行大幅改造。"
            })
        elif total_explicit_matches >= 3:
            recommendations.append({
                "level": "warning",
                "title": "理论应用优化",
                "description": f"有 {total_explicit_matches} 处理论框架使用。",
                "action": "建议在论述中增加独特的分析维度，体现差异化。"
            })
        else:
            recommendations.append({
                "level": "success",
                "title": "保持原创",
                "description": "显性理论框架使用较少。",
                "action": "继续保持当前的分析深度。"
            })

        return {
            "risk_level": risk_level,
            "risk_color": risk_color,
            "risk_description": risk_desc,
            "overlap_rate": 0, # 不再显示重复率
            "total_matches": total_matches,
            "explicit_stats": {"high": exp_high, "med": exp_med, "low": exp_low},
            "implicit_stats": {"high": imp_high, "med": imp_med, "low": imp_low},
            "top_reference_years": [{"year": y, "count": c} for y, c in top_years],
            "conflicting_cases": top_cases,
            "recommendations": recommendations
        }

    except Exception as e:
        logger.error(f"Error generating analysis summary: {e}", exc_info=True)
        return {}

def build_aggregated_report(mapping, new_items, category, db_path=DB_PATH, table_name="knowledge_points"):
    """构建聚合报告，支持三级匹配"""
    conn = None
    try:
        logger.info(f"Building report for {category} with {len(mapping)} mappings using {db_path} ({table_name})")
        conn = sqlite3.connect(db_path, timeout=30)
        c = conn.cursor()
        final_results = []
        new_item_map = {}

        # 建立 map: standard_name -> list of {context, rationale}
        for item in new_items:
            name = item['standard_name']
            if name not in new_item_map:
                new_item_map[name] = []
            new_item_map[name].append({
                "context": item['context'],
                "rationale": item.get('rationale', '') # Capture rationale
            })

        # mapping 格式: {"DB术语": {"matched_terms": [...], "match_level": "high/medium/low"}}
        for db_term, match_data in mapping.items():
            # 兼容旧格式（列表）和新格式（字典）
            if isinstance(match_data, list):
                matched_new_terms = match_data
                match_level = "high"  # 默认高度重合
            elif isinstance(match_data, dict):
                matched_new_terms = match_data.get('matched_terms', [])
                match_level = match_data.get('match_level', 'high')
            else:
                continue

            if not matched_new_terms:
                continue

            # 收集所有匹配术语的 context 和 rationale
            combined_contexts = []
            combined_rationales = []
            
            for term in matched_new_terms:
                items = new_item_map.get(term, [])
                for i in items:
                    combined_contexts.append(i['context'])
                    if i['rationale']:
                        combined_rationales.append(i['rationale'])
            
            # 选择最长的作为代表
            display_new_context = ""
            if combined_contexts:
                display_new_context = max(combined_contexts, key=len)
            
            display_rationale = ""
            if combined_rationales:
                display_rationale = max(combined_rationales, key=len)

            # 查询数据库证据
            if table_name == "theories":
                 # V2 schema: category is 'Direct Theory' or 'Implied Theory'
                 query = f'''
                    SELECT f.filename, f.publish_date, k.context
                    FROM {table_name} k
                    JOIN files f ON k.file_id = f.id
                    WHERE k.standard_name = ?
                '''
                 params = [db_term]
                 # If category is provided and NOT "All Theories", filter by it.
                 if category and category != "All Theories":
                     query += " AND k.category = ?"
                     params.append(category)
                 
                 c.execute(query, tuple(params))
            else:
                c.execute(f'''
                    SELECT f.filename, f.publish_date, k.context
                    FROM {table_name} k
                    JOIN files f ON k.file_id = f.id
                    WHERE k.standard_name = ? AND k.category = ?
                ''', (db_term, category))

            evidence_list = []
            for row in c.fetchall():
                evidence_list.append({
                    "filename": row[0],
                    "year": row[1],
                    "context": row[2]
                })

            if evidence_list:
                # 根据 match_level 设置风险等级
                risk_level_map = {
                    "high": "高度重合",
                    "medium": "次重合",
                    "low": "重合"
                }
                final_results.append({
                    "risk_level": risk_level_map.get(match_level, "高度重合"),
                    "db_term": db_term,
                    "matched_terms": matched_new_terms,
                    "new_context": display_new_context,
                    "rationale": display_rationale, # Add rationale to result
                    "evidence": evidence_list,
                    "match_level": match_level  # 添加原始等级用于排序
                })

        # 按匹配等级排序：high > medium > low
        level_priority = {"high": 0, "medium": 1, "low": 2}
        final_results.sort(key=lambda x: level_priority.get(x.get('match_level', 'high'), 0))

        logger.info(f"Built {len(final_results)} results for {category}")
        return final_results
    except Exception as e:
        logger.error(f"Error building report for {category}: {e}", exc_info=True)
        return []
    finally:
        if conn:
            conn.close()

# --- WebSocket 路由 ---

@app.websocket("/ws/{client_id}")
async def websocket_endpoint(websocket: WebSocket, client_id: str):
    # Parse query params manually since FastAPI WebSocket doesn't fully support query params in signature automatically in all versions/setups easily
    # But usually it does. Let's try to access query_params from scope if needed, or just assume it's passed.
    # Actually, we can access websocket.query_params
    mode = websocket.query_params.get("mode", "general")
    logger.info(f"New connection request from client: {client_id}, mode: {mode}")
    
    heartbeat_task = None

    try:
        await websocket.accept()
        logger.info(f"WebSocket accepted for client: {client_id}")

        # 启动心跳任务
        heartbeat_task = asyncio.create_task(send_heartbeat(websocket, interval=10))
        logger.info("Heartbeat task started")

    except Exception as e:
        logger.error(f"Failed to accept WebSocket for {client_id}: {e}", exc_info=True)
        return

    try:
        # 接收文件数据
        data = await websocket.receive_bytes()
        logger.info(f"Received {len(data) / 1024:.2f} KB from client {client_id}")

        # 步骤 1: 解析 PDF
        await websocket.send_json({
            "step": "start",
            "message": "正在解析 PDF 文档结构...",
            "progress": 10
        })
        text = extract_text_from_pdf_bytes(data)
        if not text:
            await websocket.send_json({
                "step": "error",
                "message": "无法读取 PDF 文本，请检查文件格式或大小（限制50MB）"
            })
            return

        # 步骤 2: AI 提取（关键步骤，容易超时）
        await websocket.send_json({
            "step": "extract",
            "message": "AI 正在提取关键理论与知识点（可能需要1-3分钟）...",
            "progress": 30
        })
        logger.info("Starting AI extraction phase")
        new_data = await extract_new_concepts(text, websocket)
        if not new_data:
            await websocket.send_json({
                "step": "error",
                "message": "AI 提取失败，已重试3次。请稍后再试或联系管理员"
            })
            return

        all_new_items = new_data.get('items', [])
        
        # 分离 显性提及 和 隐性逻辑
        explicit_theories = [i for i in all_new_items if i.get('mention_type') == 'explicit']
        implicit_theories = [i for i in all_new_items if i.get('mention_type') == 'implicit' or i.get('mention_type') is None] # 默认为隐性
        
        logger.info(f"Extracted {len(explicit_theories)} explicit theories and {len(implicit_theories)} implicit theories")

        # 步骤 3 & 4: 加载数据库 & AI 匹配 (根据模式不同而不同)
        await websocket.send_json({
            "step": "db_fetch",
            "message": f"正在加载 FDC 历史知识库索引 ({'精准模式' if mode == 'precision' else '普通模式'})...",
            "progress": 50
        })
        
        explicit_results = []
        implicit_results = []
        
        if mode == 'precision':
            # --- 精准查询模式 (knowledge_base_v2.db) ---
            target_db = DB_V2_PATH # Use dynamic path for V2
            target_table = "theories"
            
            # Goal 1: Explicit vs Direct Theory
            db_direct_theories = get_db_tags('Direct Theory', db_path=target_db, table_name=target_table)
            
            # Goal 2: Implicit vs All Theories (Direct + Implied)
            db_all_theories = get_db_tags(None, db_path=target_db, table_name=target_table)
            
            await websocket.send_json({
                "step": "match",
                "message": "正在进行精准双轨匹配...",
                "progress": 70
            })
            
            # Mapping
            explicit_mapping = await ai_strict_mapping(explicit_theories, db_direct_theories, 'Direct Theory')
            implicit_mapping = await ai_strict_mapping(implicit_theories, db_all_theories, 'All Theories')
            
            # Building Reports
            await websocket.send_json({
                "step": "aggregate",
                "message": "正在聚合证据链...",
                "progress": 90
            })
            
            explicit_results = build_aggregated_report(explicit_mapping, explicit_theories, 'Direct Theory', db_path=target_db, table_name=target_table)
            # For Implicit results, we matched against ALL, so we use 'All Theories' category to indicate NO filter in build_aggregated_report
            implicit_results = build_aggregated_report(implicit_mapping, implicit_theories, 'All Theories', db_path=target_db, table_name=target_table)

        else:
            # --- 普通查询模式 (knowledge_base.db - 保持原有逻辑) ---
            db_theories = get_db_tags('理论') # defaults to old DB/table

            await websocket.send_json({
                "step": "match",
                "message": "正在进行严格语义归一化匹配...",
                "progress": 70
            })
            
            explicit_mapping = await ai_strict_mapping(explicit_theories, db_theories, '理论')
            implicit_mapping = await ai_strict_mapping(implicit_theories, db_theories, '理论')

            await websocket.send_json({
                "step": "aggregate",
                "message": "正在聚合证据链...",
                "progress": 90
            })
            
            explicit_results = build_aggregated_report(explicit_mapping, explicit_theories, '理论')
            implicit_results = build_aggregated_report(implicit_mapping, implicit_theories, '理论')

        # 新增：归一化合并相同概念 (分别进行)
        await websocket.send_json({
            "step": "consolidate",
            "message": "正在归一化合并相同概念...",
            "progress": 95
        })
        logger.info("Starting consolidation phase")
        
        explicit_results = await consolidate_results(explicit_results, '理论(显性)')
        implicit_results = await consolidate_results(implicit_results, '理论(隐性)')

        # 新增：强制翻译/标准化为中文
        await websocket.send_json({
            "step": "translate",
            "message": "正在标准化中文术语...",
            "progress": 98
        })
        explicit_results = await translate_to_chinese_academic(explicit_results)
        implicit_results = await translate_to_chinese_academic(implicit_results)

        # 合并用于统计
        all_results = explicit_results + implicit_results

        # 生成智能总结
        analysis_summary = generate_analysis_summary(
            explicit_results,
            implicit_results,
            all_new_items
        )

        final_report = {
            "meta": {
                "filename": "Uploaded File",
                "detected_title": new_data.get("meta", {}).get("title", "Unknown"),
                "timestamp": datetime.now().isoformat()
            },
            "summary": {
                "high_risk_count": analysis_summary.get('total_matches', 0), 
                "medium_risk_count": 0, 
                "low_risk_count": 0,
                "total_items": len(all_new_items),
                "theory_count": len(all_new_items),
                "point_count": 0,
                "theory_match_count": len(all_results),
                "point_match_count": 0,
                "analysis": analysis_summary
            },
            "results": {
                "explicit_theories": explicit_results, # 显性
                "implicit_theories": implicit_results, # 隐性
                "theories": all_results, # 兼容旧前端
                "knowledge_points": []
            }
        }

        await websocket.send_json({
            "step": "done",
            "message": "分析完成",
            "progress": 100,
            "data": final_report
        })
        logger.info(f"Analysis completed successfully for client {client_id}")

        await asyncio.sleep(2)
        await websocket.close()

    except WebSocketDisconnect as e:
        logger.warning(f"Client {client_id} disconnected: {e}")
    except asyncio.CancelledError:
        logger.warning(f"WebSocket task cancelled for {client_id}")
    except Exception as e:
        logger.error(f"Critical error in WebSocket for {client_id}: {e}", exc_info=True)
        try:
            await websocket.send_json({
                "step": "error",
                "message": f"服务器错误: {str(e)[:100]}。请检查日志或联系管理员"
            })
            await websocket.close()
        except Exception as close_error:
            logger.error(f"Failed to send error message: {close_error}")
    finally:
        # 取消心跳任务
        if heartbeat_task and not heartbeat_task.done():
            heartbeat_task.cancel()
            try:
                await heartbeat_task
            except asyncio.CancelledError:
                pass
            logger.info(f"Heartbeat task stopped for {client_id}")

# --- 导出接口 (NEW) ---

from urllib.parse import quote

def normalize_title(title: str) -> str:
    """Normalize titles for matching."""
    if not title:
        return ""
    t = title.replace("\ufeff", "").strip().lower()
    if t.endswith(".pdf"):
        t = t[:-4]
    return t

@lru_cache(maxsize=1)
def load_case_lookup(csv_path: str = "database.csv") -> Dict[str, Dict[str, str]]:
    """Load case metadata from CSV for author/DOI enrichment."""
    lookup: Dict[str, Dict[str, str]] = {}
    path = Path(csv_path)
    if not path.exists():
        logger.warning(f"{csv_path} not found for checklist enrichment")
        return lookup
    try:
        with path.open("r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                title = row.get("案例标题") or row.get("\ufeff案例标题") or ""
                if not title:
                    continue
                norm = normalize_title(title)
                lookup[norm] = {
                    "title": title,
                    "company": row.get("公司", ""),
                    "author": row.get("第一作者", ""),
                    "doi": row.get("DOI", ""),
                }
    except Exception as e:
        logger.error(f"Failed to load {csv_path}: {e}")
    return lookup

def format_case_entry(title: str, lookup: Dict[str, Dict[str, str]]) -> str:
    """Build display string: 公司（第一作者） DOI."""
    norm = normalize_title(title)
    info = lookup.get(norm)
    if not info:
        return title
    company = info.get("company") or ""
    author = info.get("author") or ""
    doi = info.get("doi") or ""
    parts = []
    if company and author:
        parts.append(f"{company}（{author}）")
    elif company:
        parts.append(company)
    elif author:
        parts.append(author)
    if doi:
        parts.append(doi)
    return " ".join(parts) if parts else title

def build_checklist_rows(items: List[Dict[str, Any]], lookup: Dict[str, Dict[str, str]]):
    rows = []
    remark_map = {"高度重合": "高频预警", "次重合": "关注", "重合": ""}
    for idx, item in enumerate(items, 1):
        concept = item.get("db_term") or item.get("new_term") or f"项{idx}"
        evidence = item.get("evidence", [])
        seen = set()
        case_entries = []
        for ev in evidence:
            fname = ev.get("filename", "")
            norm = normalize_title(fname)
            if not fname or norm in seen:
                continue
            seen.add(norm)
            case_entries.append(format_case_entry(fname, lookup))
        count = len(case_entries) if case_entries else len(evidence)
        remark = remark_map.get(item.get("risk_level"), "")
        rows.append(
            {
                "concept": concept,
                "count": count,
                "cases": case_entries,
                "remark": remark,
            }
        )
    return rows

def generate_checklist_html_string(data: Dict[str, Any]) -> str:
    """Generate HTML in example.pdf checklist style."""
    lookup = load_case_lookup()
    results = data.get("results", {})
    explicit_rows = build_checklist_rows(results.get("explicit_theories", []), lookup)
    implicit_rows = build_checklist_rows(results.get("implicit_theories", []), lookup)

    def render_section(title: str, rows):
        if not rows:
            return ""
        body = ""
        for r in rows:
            cases = "<br>".join(r["cases"]) if r["cases"] else ""
            body += f"""
            <tr>
                <td style="border:1px solid #000;padding:6px 8px;white-space:nowrap;">{r['concept']}</td>
                <td style="border:1px solid #000;padding:6px 8px;text-align:center;white-space:nowrap;">{r['count']}</td>
                <td style="border:1px solid #000;padding:6px 8px;">{cases}</td>
                <td style="border:1px solid #000;padding:6px 8px;white-space:nowrap;">{r['remark']}</td>
            </tr>
            """
        return f"""
        <tr>
            <td colspan="4" style="padding:8px 4px;font-weight:bold;">{title}</td>
        </tr>
        {body}
        """

    return f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <title>查重报告</title>
        <style>
            body {{ font-family: "SimSun","Times New Roman",serif; margin: 1.8cm 2.2cm; color: #000; }}
            h1 {{ text-align:center; font-size:24px; margin:0; }}
            h2 {{ text-align:center; font-size:20px; margin:6px 0 16px 0; }}
            table {{ width:100%; border-collapse: collapse; font-size:12px; }}
            .meta td {{ padding:4px 6px; }}
            .section-title {{ font-weight:bold; margin-top:12px; }}
        </style>
    </head>
    <body>
        <h1>复旦管理案例库</h1>
        <h2>投稿案例教学参考知识点查重报告</h2>
        <table class="meta">
            <tr><td>1. 稿件来源：</td><td></td></tr>
            <tr><td>2. 投稿编号：</td><td></td></tr>
            <tr><td>3. 案例名称：</td><td></td></tr>
            <tr><td>4. 查重时间：</td><td></td></tr>
            <tr><td>5. 查重结果：</td><td></td></tr>
        </table>
        <p style="margin:10px 0 6px 0;font-weight:bold;">5. 查重结果：</p>
        <table style="border:1px solid #000; border-collapse: collapse;">
            <tr>
                <th style="border:1px solid #000;padding:6px 8px;white-space:nowrap;">理论/工具/模型</th>
                <th style="border:1px solid #000;padding:6px 8px;white-space:nowrap;">重合数</th>
                <th style="border:1px solid #000;padding:6px 8px;">已入库案例企业、一作及 DOI 编号</th>
                <th style="border:1px solid #000;padding:6px 8px;white-space:nowrap;">备注</th>
            </tr>
            {render_section("主概念（显性提及）", explicit_rows)}
            {render_section("次概念（隐性提及）", implicit_rows)}
        </table>
        <p style="font-size:10px;margin-top:12px;">复旦管理案例库以教学应用为导向，为更好地满足教学需要，提高入库案例产品质量，对所有投稿案例采用严谨的学术评审流程。案例中心通过动态搭建最新五年入库案例知识图谱并借助人工智能工具，对新投稿案例的知识点进行查重，供同行评议参考。</p>
    </body>
    </html>
    """

def generate_checklist_docx_bytes(data: Dict[str, Any]) -> bytes:
    """Generate Word docx for checklist style."""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        logger.error(f"python-docx not available: {e}")
        raise

    lookup = load_case_lookup()
    results = data.get("results", {})
    explicit_rows = build_checklist_rows(results.get("explicit_theories", []), lookup)
    implicit_rows = build_checklist_rows(results.get("implicit_theories", []), lookup)

    def add_section(doc: Any, title: str, rows):
        if not rows:
            return
        doc.add_paragraph(title).runs[0].bold = True
        table = doc.add_table(rows=len(rows) + 1, cols=4)
        table.style = "Table Grid"
        headers = ["理论/工具/模型", "重合数", "已入库案例企业、一作及 DOI 编号", "备注"]
        for i, h in enumerate(headers):
            run = table.rows[0].cells[i].paragraphs[0].add_run(h)
            run.bold = True
        for ridx, r in enumerate(rows, 1):
            table.rows[ridx].cells[0].text = str(r["concept"])
            table.rows[ridx].cells[1].text = str(r["count"])
            table.rows[ridx].cells[2].text = "\n".join(r["cases"])
            table.rows[ridx].cells[3].text = r["remark"]

    doc = Document()
    title1 = doc.add_heading("复旦管理案例库", level=1)
    title1.alignment = 1
    title2 = doc.add_heading("投稿案例教学参考知识点查重报告", level=2)
    title2.alignment = 1

    meta = [
        "1. 稿件来源：",
        "2. 投稿编号：",
        "3. 案例名称：",
        "4. 查重时间：",
        "5. 查重结果：",
    ]
    for line in meta:
        doc.add_paragraph(line)

    add_section(doc, "主概念（显性提及）", explicit_rows)
    add_section(doc, "次概念（隐性提及）", implicit_rows)

    note = (
        "复旦管理案例库以教学应用为导向，为更好地满足教学需要，提高入库案例产品质量，"
        "对所有投稿案例采用严谨的学术评审流程。案例中心通过动态搭建最新五年入库案例知识图谱并借助人工智能工具，"
        "对新投稿案例的知识点进行查重，供同行评议参考。"
    )
    doc.add_paragraph(note).runs[0].font.size = Pt(9)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

@app.post("/api/export")
async def export_report(report: ReportData):
    """
    接收前端传来的 JSON 报告数据，生成美化后的 HTML 并返回
    """
    print("Export request received...", flush=True)
    
    report_data = report.model_dump()
    report_type = report.report_type

    if report_type == "paper":
        html_content = generate_paper_html_string(report_data)
        raw_filename = f"FDC_Paper_Report_{report.meta.get('filename', 'report').replace('.pdf', '')}.html"
        encoded_filename = quote(raw_filename)
        return HTMLResponse(
            content=html_content, 
            media_type="text/html", 
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    elif report_type == "checklist":
        html_content = generate_checklist_html_string(report_data)
        raw_filename = f"FDC_Checklist_Report_{report.meta.get('filename', 'report').replace('.pdf', '')}.html"
        encoded_filename = quote(raw_filename)
        return HTMLResponse(
            content=html_content, 
            media_type="text/html", 
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    elif report_type == "checklist_word":
        docx_bytes = generate_checklist_docx_bytes(report_data)
        raw_filename = f"FDC_Checklist_Report_{report.meta.get('filename', 'report').replace('.pdf', '')}.docx"
        encoded_filename = quote(raw_filename)
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    else:
        html_content = generate_html_string(report_data)
        raw_filename = f"FDC_Dashboard_Report_{report.meta.get('filename', 'report').replace('.pdf', '')}.html"
        encoded_filename = quote(raw_filename)
        return HTMLResponse(
            content=html_content, 
            media_type="text/html", 
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )

def generate_html_string(data):
    # 将 JSON 转换为高级 HTML 字符串
    
    summary = data.get('summary', {})
    analysis = summary.get('analysis', {})
    
    # --- 1. SVG Chart Generation (Simple Bar Chart) ---
    def generate_svg_chart(year_data):
        if not year_data:
            return ""
        
        # Sort by year
        sorted_data = sorted(year_data, key=lambda x: x['year'])
        max_val = max([x['count'] for x in sorted_data]) if sorted_data else 1
        
        svg_width = 600
        svg_height = 200
        bar_width = 40
        gap = 30
        
        bars_svg = ""
        total_width = len(sorted_data) * (bar_width + gap)
        start_x = (svg_width - total_width) / 2
        
        for i, item in enumerate(sorted_data):
            x = start_x + i * (bar_width + gap)
            bar_h = (item['count'] / max_val) * (svg_height - 40)
            y = svg_height - 20 - bar_h
            
            bars_svg += f"""
            <rect x="{x}" y="{y}" width="{bar_width}" height="{bar_h}" fill="#334155" rx="2" />
            <text x="{x + bar_width/2}" y="{svg_height}" text-anchor="middle" font-family="monospace" font-size="12" fill="#64748b">{item['year']}</text>
            <text x="{x + bar_width/2}" y="{y - 5}" text-anchor="middle" font-family="sans-serif" font-size="12" font-weight="bold" fill="#334155">{item['count']}</text>
            """
            
        return f"""
        <svg width="100%" height="200" viewBox="0 0 {svg_width} {svg_height}" xmlns="http://www.w3.org/2000/svg">
            <line x1="0" y1="{svg_height-20}" x2="{svg_width}" y2="{svg_height-20}" stroke="#e2e8f0" stroke-width="1" />
            {bars_svg}
        </svg>
        """

    chart_html = generate_svg_chart(analysis.get('top_reference_years', []))

    def render_items(items):
        if not items:
            return '<div class="text-center py-12 text-slate-400 font-serif italic">表现优秀。未检测到理论框架重复。</div>'
        
        html = '<div class="flex flex-col gap-8">'
        for item in items:
            evidence_html = ""
            for ev in item['evidence']:
                evidence_html += f"""
                <div class="mb-6 last:mb-0 group">
                    <div class="flex items-baseline justify-between mb-2">
                        <span class="text-sm font-bold text-slate-800 border-b border-transparent group-hover:border-slate-800 transition-all">{ev['filename']}</span>
                        <span class="text-xs font-mono text-slate-500">{ev['year']}</span>
                    </div>
                    <p class="text-sm text-slate-600 leading-relaxed font-serif pl-4 border-l border-slate-200">
                        {ev['context'][:250]}...
                    </p>
                </div>
                """
            
            variants = ", ".join(item['matched_terms']) if item.get('matched_terms') else item.get('new_term', '')
            rationale = item.get('rationale', 'AI未提供详细理由')

            # Risk Styling
            risk_level = item.get('risk_level', '高度重合')
            
            # 关键修改：默认都不展开 (details_open 为空)
            details_open = '' 
            summary_hint = '<span class="ml-4 text-xs text-slate-400 italic font-normal">(点击展开详细证据链)</span>'

            if risk_level == '高度重合': # > 5
                risk_badge = '<span class="inline-block px-2 py-1 bg-black text-white text-xs font-bold uppercase tracking-widest">高度重合</span>'
                border_style = 'border-l-4 border-black'
                match_desc = '历史库中出现 5 次以上'
            elif risk_level == '次重合': # 3-4
                risk_badge = '<span class="inline-block px-2 py-1 bg-slate-200 text-slate-800 text-xs font-bold uppercase tracking-widest">次重合</span>'
                border_style = 'border-l-4 border-slate-300'
                match_desc = '历史库中出现 3-4 次'
            else:  # 重合 1-2
                risk_badge = '<span class="inline-block px-2 py-1 bg-slate-50 text-slate-500 text-xs font-bold uppercase tracking-widest">重合</span>'
                border_style = 'border-l-4 border-slate-100'
                match_desc = '历史库中出现 1-2 次'

            # 使用 <details> 标签实现折叠/展开
            html += f"""
            <div class="break-inside-avoid mb-8">
                <details {details_open} class="group">
                    <summary class="list-none cursor-pointer">
                         <div class="flex items-start gap-6">
                            <div class="flex-1 {border_style} pl-6">
                                <div class="flex items-center gap-4 mb-3">
                                    {risk_badge}
                                    <h3 class="text-xl font-serif font-bold text-slate-900">{item['db_term']}</h3>
                                    <span class="text-xs text-slate-400 font-mono border border-slate-200 px-2 py-0.5 rounded">{match_desc}</span>
                                    {summary_hint}
                                </div>
                            </div>
                        </div>
                    </summary>
                    
                    <div class="pl-6 ml-6 mt-4 border-l border-slate-100">
                        <div class="text-sm text-slate-500 font-mono mb-6">相关变体: {variants}</div>
                        <div class="grid grid-cols-1 lg:grid-cols-2 gap-12">
                            <div>
                                <h4 class="text-xs font-bold text-slate-400 uppercase tracking-widest mb-4">当前文稿语境</h4>
                                <p class="text-base text-slate-800 leading-loose font-serif text-justify mb-6">
                                    {item['new_context']}
                                </p>
                                
                                <!-- AI Annotation Card -->
                                <div class="p-4 bg-slate-100/50 rounded-sm border border-slate-200">
                                    <h4 class="text-xs font-bold text-slate-500 uppercase tracking-widest mb-2 flex items-center gap-2">
                                        AI 判定理由
                                    </h4>
                                    <p class="text-sm text-slate-600 leading-relaxed font-serif">
                                        {rationale}
                                    </p>
                                </div>
                            </div>
                            <div>
                                <h4 class="text-xs font-bold text-slate-400 uppercase tracking-widest mb-4">历史库证据链 ({len(item['evidence'])})</h4>
                                <div class="bg-slate-50 p-6 rounded-sm">
                                    {evidence_html}
                                </div>
                            </div>
                        </div>
                    </div>
                </details>
            </div>
            """
        html += "</div>"
        return html

    # Summary Section
    risk_level = analysis.get('risk_level', '未知')
    risk_desc = analysis.get('risk_description', '')
    
    # Determine Header Color
    header_color = "text-slate-900"
    if risk_level == "高风险":
        header_color = "text-rose-700"
    elif risk_level == "逻辑风险":
        header_color = "text-amber-700"
    elif risk_level == "中度风险":
        header_color = "text-orange-600"
    elif risk_level == "良好":
        header_color = "text-emerald-700"

    # Explicit/Implicit Stats
    exp_stats = analysis.get('explicit_stats', {})
    imp_stats = analysis.get('implicit_stats', {})

    return f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <title>FDC Report - {data['meta']['detected_title']}</title>
        <script src="https://cdn.tailwindcss.com"></script>
        <link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;700&family=Inter:wght@300;400;600&display=swap" rel="stylesheet">
        <style>
            @media print {{
                body {{ -webkit-print-color-adjust: exact; }}
                .no-print {{ display: none; }}
                .page-break {{ page-break-before: always; }}
            }}
            body {{ font-family: 'Inter', sans-serif; background-color: #fff; color: #0F172A; }}
            .font-serif {{ font-family: 'Noto Serif SC', serif; }}
            details > summary {{ list-style: none; }}
            details > summary::-webkit-details-marker {{ display: none; }}
        </style>
    </head>
    <body class="p-16 max-w-6xl mx-auto">
        
        <!-- Header -->
        <header class="mb-20">
            <div class="flex justify-between items-end border-b border-slate-900 pb-8">
                <div>
                    <h1 class="text-5xl font-serif font-bold text-slate-900 mb-4 tracking-tight">商业案例理论查重报告</h1>
                    <div class="text-xl text-slate-500 font-serif italic">{data['meta']['detected_title']}</div>
                </div>
                <div class="text-right">
                    <div class="text-xs font-bold uppercase tracking-widest text-slate-400 mb-1">生成日期</div>
                    <div class="text-sm font-mono text-slate-900">{datetime.now().strftime('%Y.%m.%d')}</div>
                </div>
            </div>
        </header>

        <!-- Executive Summary -->
        <section class="mb-24">
            <h2 class="text-xs font-bold uppercase tracking-widest text-slate-400 mb-8">分析结论摘要</h2>
            
            <div class="grid grid-cols-1 md:grid-cols-2 gap-16 items-start">
                <div>
                    <div class="flex items-baseline gap-4 mb-6">
                        <span class="text-6xl font-serif font-bold {header_color}">{risk_level}</span>
                    </div>
                    <p class="text-lg text-slate-700 leading-relaxed font-serif mb-8">
                        {risk_desc}
                    </p>
                    
                    <div class="grid grid-cols-2 gap-8 border-t border-slate-100 pt-8">
                        <div>
                             <h4 class="text-xs font-bold uppercase tracking-widest text-slate-400 mb-4">显性框架 (Explicit)</h4>
                             <div class="flex gap-4">
                                <div>
                                    <div class="text-2xl font-bold text-slate-900">{exp_stats.get('high', 0)}</div>
                                    <div class="text-[10px] text-slate-400">高度</div>
                                </div>
                                <div>
                                    <div class="text-2xl font-bold text-slate-900">{exp_stats.get('med', 0)}</div>
                                    <div class="text-[10px] text-slate-400">次重</div>
                                </div>
                             </div>
                        </div>
                        <div>
                             <h4 class="text-xs font-bold uppercase tracking-widest text-slate-400 mb-4">隐性逻辑 (Implicit)</h4>
                             <div class="flex gap-4">
                                <div>
                                    <div class="text-2xl font-bold text-slate-900">{imp_stats.get('high', 0)}</div>
                                    <div class="text-[10px] text-slate-400">高度</div>
                                </div>
                                <div>
                                    <div class="text-2xl font-bold text-slate-900">{imp_stats.get('med', 0)}</div>
                                    <div class="text-[10px] text-slate-400">次重</div>
                                </div>
                             </div>
                        </div>
                    </div>
                </div>
                
                <!-- Chart Area -->
                <div class="bg-slate-50 p-8 rounded-sm">
                    <h3 class="text-xs font-bold uppercase tracking-widest text-slate-400 mb-6 text-center">历史引用趋势分布</h3>
                    {chart_html}
                </div>
            </div>
        </section>

        <div class="page-break"></div>

        <!-- Detailed Findings -->
        <section>
            <div class="flex items-center justify-between mb-8">
                 <h2 class="text-xl font-bold uppercase tracking-widest text-slate-800 border-l-4 border-slate-900 pl-4">显性提及 (Explicit Mentions)</h2>
                 <span class="text-sm text-slate-400 font-mono">文中使用标准理论名称的部分</span>
            </div>
            
            {render_items(data['results']['explicit_theories'])}
            
            <div class="h-24"></div> <!-- Spacer -->

            <div class="flex items-center justify-between mb-8">
                 <h2 class="text-xl font-bold uppercase tracking-widest text-slate-800 border-l-4 border-slate-400 pl-4">隐性逻辑 (Implicit Logic)</h2>
                 <span class="text-sm text-slate-400 font-mono">文中运用理论逻辑但未点名的部分</span>
            </div>
            
            {render_items(data['results']['implicit_theories'])}
        </section>

        <!-- Footer -->
        <footer class="mt-32 pt-8 border-t border-slate-100 text-center">
             <div class="text-xs text-slate-300 font-mono uppercase tracking-widest">FDC 智能案例评审系统</div>
        </footer>

        <script>
            window.onload = function() {{ setTimeout(function() {{ window.print(); }}, 800); }}
        </script>
    </body>
    </html>
    """

def generate_paper_html_string(data):
    """
    Generates an HTML string for an academic paper-style report.
    """
    summary = data.get('summary', {})
    analysis = summary.get('analysis', {})

    # Helper to render items in a paper style (no collapse)
    def render_paper_items(items, section_title):
        if not items:
            return f"""<h3 style="font-family: 'Noto Serif SC', serif; font-size: 1.25em; margin-bottom: 0.5em;">{section_title}</h3><p style="font-family: 'Noto Serif SC', serif; font-size: 0.9em; line-height: 1.6;">未检测到相关理论。</p>"""

        html = f"""<h3 style="font-family: 'Noto Serif SC', serif; font-size: 1.25em; margin-bottom: 0.5em;">{section_title}</h3>\n"""
        for idx, item in enumerate(items):
            html += f"""
            <div style="margin-bottom: 2em; border-bottom: 1px dashed #eee; padding-bottom: 1em;">
                <p style="font-family: 'Noto Serif SC', serif; font-weight: bold; font-size: 1.1em; margin-bottom: 0.5em;">{idx + 1}. {item['db_term']} (风险等级: {item['risk_level']})</p>
                <p style="font-family: 'Noto Serif SC', serif; font-size: 0.95em; line-height: 1.6; margin-bottom: 0.5em;"><b>当前文稿语境：</b>{item['new_context']}</p>
                <p style="font-family: 'Noto Serif SC', serif; font-size: 0.9em; line-height: 1.5; color: #555; margin-bottom: 1em;"><b>AI判定理由：</b>{item.get('rationale', 'AI未提供详细理由')}</p>
                <p style="font-family: 'Noto Serif SC', serif; font-size: 0.95em; line-height: 1.6; margin-bottom: 0.5em;"><b>相关变体：</b>{', '.join(item.get('matched_terms', []))}</p>
                <p style="font-family: 'Noto Serif SC', serif; font-weight: bold; font-size: 0.95em; margin-top: 1em;">历史库证据链 ({len(item['evidence'])})：</p>
                <ul style="list-style-type: disc; padding-left: 1.5em; margin-top: 0.5em;">
            """
            for ev_idx, ev in enumerate(item['evidence']):
                html += f"""
                    <li style="font-family: 'Noto Serif SC', serif; font-size: 0.85em; line-height: 1.5; margin-bottom: 0.5em;">
                        <b>{ev['filename']} ({ev['year']})</b>: {ev['context']}
                    </li>
                """
            html += f"""
                </ul>
            </div>
            """
        html += "\n"
        return html

    # Generate academic style table for conflicting cases
    def generate_paper_case_table(case_data):
        if not case_data:
            return "<p style=\"font-family: 'Noto Serif SC', serif; font-size: 0.9em; line-height: 1.6;\">未发现显著的历史案例重复。</p>"
        
        table_rows = ""
        for idx, item in enumerate(case_data):
            table_rows += f"""
            <tr>
                <td style="padding: 8px 12px; border: 1px solid #000; text-align: center;">{idx + 1}</td>
                <td style="padding: 8px 12px; border: 1px solid #000;">{item['filename'].replace('.pdf', '')}</td>
                <td style="padding: 8px 12px; border: 1px solid #000; text-align: center;">{item['year']}</td>
                <td style="padding: 8px 12px; border: 1px solid #000; text-align: center;">{item['count']}</td>
            </tr>
            """

        return f"""
        <table style="width: 100%; border-collapse: collapse; margin-top: 1em; font-family: 'Noto Serif SC', serif; font-size: 0.9em;">
            <caption>表1: 重复度最高的历史案例</caption>
            <thead>
                <tr>
                    <th style="padding: 8px 12px; border: 1px solid #000; background-color: #f2f2f2; text-align: center;">序号</th>
                    <th style="padding: 8px 12px; border: 1px solid #000; background-color: #f2f2f2; text-align: center;">案例标题</th>
                    <th style="padding: 8px 12px; border: 1px solid #000; background-color: #f2f2f2; text-align: center;">年份</th>
                    <th style="padding: 8px 12px; border: 1px solid #000; background-color: #f2f2f2; text-align: center;">重复点数</th>
                </tr>
            </thead>
            <tbody>
                {table_rows}
            </tbody>
        </table>
        """

    # Generate academic style table for year trends
    def generate_paper_year_table(year_data):
        if not year_data:
            return "<p style=\"font-family: 'Noto Serif SC', serif; font-size: 0.9em; line-height: 1.6;\">暂无历史引用数据。</p>"

        sorted_data = sorted(year_data, key=lambda x: x['year'])
        table_rows = ""
        for item in sorted_data:
            table_rows += f"""
            <tr>
                <td style="padding: 8px 12px; border: 1px solid #000; text-align: center;">{item['year']}</td>
                <td style="padding: 8px 12px; border: 1px solid #000; text-align: center;">{item['count']}</td>
            </tr>
            """
        return f"""
        <table style="width: 50%; border-collapse: collapse; margin-top: 1em; font-family: 'Noto Serif SC', serif; font-size: 0.9em; margin-left: auto; margin-right: auto;">
            <caption>表2: 历史引用趋势分布</caption>
            <thead>
                <tr>
                    <th style="padding: 8px 12px; border: 1px solid #000; background-color: #f2f2f2; text-align: center;">年份</th>
                    <th style="padding: 8px 12px; border: 1px solid #000; background-color: #f2f2f2; text-align: center;">引用次数</th>
                </tr>
            </thead>
            <tbody>
                {table_rows}
            </tbody>
        </table>
        """

    title = data['meta'].get('detected_title', '商业案例理论查重报告')
    current_date = datetime.now().strftime('%Y年%m月%d日 %H:%M')
    risk_level = analysis.get('risk_level', '未知')
    risk_desc = analysis.get('risk_description', '系统正在分析文稿的原创性...')
    overlap_rate = analysis.get('overlap_rate', 0)
    exp_stats = analysis.get('explicit_stats', {})
    imp_stats = analysis.get('implicit_stats', {})
    recommendations_html = ""
    for rec in analysis.get('recommendations', []):
        recommendations_html += f"""
        <p style="font-family: 'Noto Serif SC', serif; font-size: 0.95em; line-height: 1.6; margin-bottom: 0.5em;">
            <b>{rec['title']}：</b> {rec['description']} {rec['action']}
        </p>
        """

    return f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <title>{title} - 文字版报告</title>
        <link href="https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@300;400;700&family=Inter:wght@300;400;600&display=swap" rel="stylesheet">
        <style>
            body {{ font-family: 'Noto Serif SC', serif; line-height: 1.8; margin: 2cm 3cm; font-size: 11pt; color: #333; }}
            h1, h2, h3, h4 {{ font-family: 'Noto Serif SC', serif; font-weight: bold; margin-top: 1.5em; margin-bottom: 0.8em; line-height: 1.3; }}
            h1 {{ font-size: 2em; text-align: center; margin-bottom: 1em; }}
            h2 {{ font-size: 1.6em; border-bottom: 1px solid #ccc; padding-bottom: 0.3em; margin-top: 2em; }}
            h3 {{ font-size: 1.25em; margin-top: 1.5em; }}
            p {{ margin-bottom: 1em; }}
            b {{ font-weight: bold; }}
            table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
            th, td {{ border: 1px solid #000; padding: 8px 12px; text-align: left; }}
            th {{ background-color: #f2f2f2; font-weight: bold; }}
            caption {{ caption-side: top; text-align: center; margin-bottom: 0.5em; font-weight: bold; font-size: 0.95em; }}
            .page-break {{ page-break-before: always; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        <p style="text-align: right; margin-bottom: 2em;">生成日期：{current_date}</p>

        <h2>1. 综合分析结论</h2>
        <p>
            本报告对文稿 <b>“{title}”</b> 进行了理论框架重合度分析。
            综合评估结果显示，文稿的原创度风险等级为：<b>{risk_level}</b>。
            具体分析如下：
        </p>
        <h3>1.1 风险评估概览</h3>
        <p>{risk_desc}</p>
        
        <h3>1.2 显性理论框架统计</h3>
        <p>
            在文稿中<b>显性提及</b>的理论框架中，存在 <b>{exp_stats.get('high', 0)}</b> 处高度重合，
            <b>{exp_stats.get('med', 0)}</b> 处次重合，以及 <b>{exp_stats.get('low', 0)}</b> 处一般重合。
        </p>

        <h3>1.3 隐性逻辑结构统计</h3>
        <p>
            在文稿中<b>隐性逻辑结构</b>方面，存在 <b>{imp_stats.get('high', 0)}</b> 处高度重合，
            <b>{imp_stats.get('med', 0)}</b> 处次重合，以及 <b>{imp_stats.get('low', 0)}</b> 处一般重合。
        </p>
        
        <h3>1.4 修改建议</h3>
        {recommendations_html}

        <div class="page-break"></div>

        <h2>2. 理论框架匹配详情</h2>
        <h3>2.1 显性提及理论</h3>
        {render_paper_items(data['results']['explicit_theories'], '显性提及理论列表')}

        <div class="page-break"></div>

        <h3>2.2 隐性逻辑理论</h3>
        {render_paper_items(data['results']['implicit_theories'], '隐性逻辑理论列表')}

        <div class="page-break"></div>

        <h2>3. 历史引用分析</h2>
        <h3>3.1 历史引用趋势分布</h3>
        {generate_paper_year_table(analysis.get('top_reference_years', []))}
        <p style="text-align: center; font-size: 0.9em; margin-top: 0.5em;">注：上表展示了各年份历史案例中相关理论的引用次数。</p>

        <h3>3.2 重复度最高的历史案例</h3>
        {generate_paper_case_table(analysis.get('conflicting_cases', []))}
        <p style="text-align: center; font-size: 0.9em; margin-top: 0.5em;">注：上表列出了与当前文稿理论框架重合点数最多的前10个历史案例。</p>

        <footer style="text-align: center; margin-top: 3em; padding-top: 1em; border-top: 1px solid #eee; font-size: 0.8em; color: #777;">
            FDC 智能案例评审系统 - 本报告由AI自动生成，仅供参考。
        </footer>
    </body>
    </html>
    """

# --- 辅助函数 ---

async def send_heartbeat(websocket: WebSocket, interval: int = 10):
    """WebSocket 心跳任务，定期发送ping保持连接活跃"""
    try:
        while True:
            await asyncio.sleep(interval)
            try:
                await websocket.send_json({"type": "heartbeat", "timestamp": datetime.now().isoformat()})
            except Exception as e:
                logger.warning(f"Heartbeat failed: {e}")
                break
    except asyncio.CancelledError:
        logger.info("Heartbeat task cancelled")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app,
        host="0.0.0.0", # 监听所有网络接口，允许局域网访问
        port=8000,
        timeout_keep_alive=600,  # 增加到600秒，支持长时间AI处理
        ws_ping_interval=30,      # WebSocket ping间隔30秒
        ws_ping_timeout=60        # WebSocket ping超时60秒
    )
